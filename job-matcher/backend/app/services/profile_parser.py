"""Parse CV (PDF/DOCX) and LinkedIn data export into structured profile data."""

import os
import json
import csv
import zipfile
from typing import Dict, Any, Optional
from datetime import datetime

import pdfplumber
from docx import Document

from app.services.llm_service import parse_cv_with_llm


async def parse_cv(filepath: str, ext: str) -> Dict[str, Any]:
    """Parse a CV file and extract structured profile data."""
    if ext == "pdf":
        text = _extract_pdf_text(filepath)
    elif ext == "docx":
        text = _extract_docx_text(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    parsed = await parse_cv_with_llm(text)
    return parsed


def _extract_pdf_text(filepath: str) -> str:
    """Extract text from a PDF file."""
    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _extract_docx_text(filepath: str) -> str:
    """Extract text from a DOCX file."""
    doc = Document(filepath)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    return "\n".join(text_parts)


async def parse_linkedin_export(zip_filepath: str) -> Dict[str, Any]:
    """Parse a LinkedIn data export ZIP file."""
    profile_data = {}

    with zipfile.ZipFile(zip_filepath, "r") as z:
        file_list = z.namelist()

        # Parse Profile.csv
        profile_csv = _find_file(file_list, "Profile.csv")
        if profile_csv:
            with z.open(profile_csv) as f:
                content = f.read().decode("utf-8")
                reader = csv.DictReader(content.splitlines())
                for row in reader:
                    profile_data["name"] = f"{row.get('First Name', '')} {row.get('Last Name', '')}".strip()
                    profile_data["headline"] = row.get("Headline", "")
                    profile_data["summary"] = row.get("Summary", "")
                    profile_data["location"] = row.get("Geo Location", "")
                    profile_data["email"] = row.get("Email Address", "")
                    break

        # Parse Positions.csv (work experience)
        positions_csv = _find_file(file_list, "Positions.csv")
        if positions_csv:
            with z.open(positions_csv) as f:
                content = f.read().decode("utf-8")
                reader = csv.DictReader(content.splitlines())
                experience = []
                for row in reader:
                    exp = {
                        "title": row.get("Title", ""),
                        "company": row.get("Company Name", ""),
                        "description": row.get("Description", ""),
                        "location": row.get("Location", ""),
                        "started_on": row.get("Started On", ""),
                        "finished_on": row.get("Finished On", ""),
                    }
                    experience.append(exp)
                profile_data["experience"] = experience
                if experience:
                    profile_data["years_experience"] = _estimate_years(experience)

        # Parse Skills.csv
        skills_csv = _find_file(file_list, "Skills.csv")
        if skills_csv:
            with z.open(skills_csv) as f:
                content = f.read().decode("utf-8")
                reader = csv.DictReader(content.splitlines())
                skills = [row.get("Name", "") for row in reader if row.get("Name")]
                profile_data["skills"] = skills

        # Parse Education.csv
        education_csv = _find_file(file_list, "Education.csv")
        if education_csv:
            with z.open(education_csv) as f:
                content = f.read().decode("utf-8")
                reader = csv.DictReader(content.splitlines())
                education = []
                for row in reader:
                    edu = {
                        "school": row.get("School Name", ""),
                        "degree": row.get("Degree Name", ""),
                        "field": row.get("Notes", "") or row.get("Field of Study", ""),
                        "start_date": row.get("Start Date", ""),
                        "end_date": row.get("End Date", ""),
                    }
                    education.append(edu)
                profile_data["education"] = education

        # Parse Languages.csv
        languages_csv = _find_file(file_list, "Languages.csv")
        if languages_csv:
            with z.open(languages_csv) as f:
                content = f.read().decode("utf-8")
                reader = csv.DictReader(content.splitlines())
                languages = [row.get("Name", "") for row in reader if row.get("Name")]
                profile_data["languages"] = languages

        # Parse Certifications.csv
        certs_csv = _find_file(file_list, "Certifications.csv")
        if certs_csv:
            with z.open(certs_csv) as f:
                content = f.read().decode("utf-8")
                reader = csv.DictReader(content.splitlines())
                certifications = []
                for row in reader:
                    cert = {
                        "name": row.get("Name", ""),
                        "authority": row.get("Authority", ""),
                        "started_on": row.get("Started On", ""),
                    }
                    certifications.append(cert)
                profile_data["certifications"] = certifications

    return profile_data


def _find_file(file_list: list, filename: str) -> Optional[str]:
    for f in file_list:
        if f.endswith(filename) or f.endswith(f"/{filename}"):
            return f
    return None


def _estimate_years(experience: list) -> int:
    total_months = 0
    for exp in experience:
        start = exp.get("started_on", "")
        end = exp.get("finished_on", "")
        if start:
            try:
                start_date = _parse_date(start)
                end_date = _parse_date(end) if end else datetime.now()
                months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
                total_months += max(0, months)
            except (ValueError, TypeError):
                continue
    return max(1, total_months // 12)


def _parse_date(date_str: str) -> datetime:
    from dateutil.parser import parse
    return parse(date_str)
